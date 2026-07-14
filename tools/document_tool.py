import os
from datetime import datetime
from docx import Document
from docx.shared import Pt


class DocumentTool:

    OUTPUT_DIR = "output"

    def __init__(self):

        os.makedirs(self.OUTPUT_DIR, exist_ok=True)

    def create_document(self, plan: dict, sections: dict):

        document = Document()

        title = document.add_heading(
            plan["document_title"],
            level=1
        )

        title.runs[0].font.size = Pt(22)

        document.add_paragraph(
            f"Document Type : {plan['document_type']}"
        )

        document.add_paragraph(
            f"Generated On : {datetime.now()}"
        )

        if plan["assumptions"]:

            document.add_heading(
                "Assumptions",
                level=2
            )

            for assumption in plan["assumptions"]:

                document.add_paragraph(
                    assumption,
                    style="List Bullet"
                )

        for heading, content in sections.items():

            document.add_heading(
                heading,
                level=2
            )

            document.add_paragraph(content)

        filename = (
            plan["document_title"]
            .replace(" ", "_")
            + ".docx"
        )

        filepath = os.path.join(
            self.OUTPUT_DIR,
            filename
        )

        document.save(filepath)

        return filepath


document_tool = DocumentTool()